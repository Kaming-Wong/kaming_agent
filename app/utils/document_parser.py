import logging
import os
import tempfile
from typing import List, Optional
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

logger = logging.getLogger(__name__)

# 支持的文件类型
SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".doc"}


class DocumentParser:
    """文档解析器 — 解析上传文件并切分为 chunks"""

    CHUNK_SIZE = 800
    CHUNK_OVERLAP = 100

    def __init__(self, chunk_size: int = CHUNK_SIZE, chunk_overlap: int = CHUNK_OVERLAP):
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""],
        )

    def parse(self, file_path: str, filename: str) -> str:
        """解析文件为纯文本

        Args:
            file_path: 本地文件路径
            filename: 原文件名（用于判断类型）
        Returns:
            提取的文本内容
        """
        ext = os.path.splitext(filename)[1].lower()

        if ext == ".txt":
            return self._parse_txt(file_path)
        elif ext == ".md":
            return self._parse_txt(file_path)  # Markdown 按纯文本处理
        elif ext == ".pdf":
            return self._parse_pdf(file_path)
        elif ext in (".docx", ".doc"):
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {ext}，支持: {', '.join(SUPPORTED_EXTENSIONS)}")

    def parse_and_chunk(self, file_path: str, filename: str,
                        source_info: Optional[dict] = None) -> List[Document]:
        """解析文件并切分成 chunks

        Args:
            file_path: 本地文件路径
            filename: 原文件名
            source_info: 附加元数据（如 session_id）
        Returns:
            Document 列表，每个 Document 包含 page_content 和 metadata
        """
        text = self.parse(file_path, filename)
        if not text.strip():
            logger.warning(f"Empty text extracted from {filename}")
            return []

        metadata = {"source": filename, **(source_info or {})}
        documents = [Document(page_content=text, metadata=metadata)]
        chunks = self.splitter.split_documents(documents)

        # 给每个 chunk 添加序号
        for i, chunk in enumerate(chunks):
            chunk.metadata["chunk_index"] = i
            chunk.metadata["chunk_count"] = len(chunks)

        logger.info(f"Parsed {filename}: {len(text)} chars → {len(chunks)} chunks (size={self.CHUNK_SIZE}, overlap={self.CHUNK_OVERLAP})")
        return chunks

    # ── 各格式解析 ──

    @staticmethod
    def _parse_txt(file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            texts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    texts.append(text)
            return "\n".join(texts)
        except ImportError:
            logger.warning("pypdf not installed, trying pdfminer...")
            return DocumentParser._parse_pdf_fallback(file_path)

    @staticmethod
    def _parse_pdf_fallback(file_path: str) -> str:
        """pdfminer 兜底"""
        try:
            from pdfminer.high_level import extract_text
            return extract_text(file_path)
        except ImportError:
            raise ImportError("请安装 pypdf 或 pdfminer.six: pip install pypdf")

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            texts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    texts.append(para.text)
            return "\n".join(texts)
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

    @classmethod
    def is_supported(cls, filename: str) -> bool:
        ext = os.path.splitext(filename)[1].lower()
        return ext in SUPPORTED_EXTENSIONS
